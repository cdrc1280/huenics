import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding in dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def create_document():
    doc = Document()
    
    # Page Setup: Standard Letter with 0.8-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Default typography
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    
    # Document Header
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_hdr_badge = p_hdr.add_run("ENTERPRISE TECHNICAL & COMMERCIAL PROPOSAL\n")
    r_hdr_badge.font.size = Pt(9)
    r_hdr_badge.font.bold = True
    r_hdr_badge.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    r_title = p_hdr.add_run("HUENICS SALES, INGESTION, INVENTORY &\nCUSTOMER QUOTATION SYSTEM\n")
    r_title.font.size = Pt(18)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    r_sub = p_hdr.add_run("Technical Scope Optimization & Modified Budget Proposal (All-In Package Included)\n")
    r_sub.font.size = Pt(11)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    
    # Metadata Table
    meta_table = doc.add_table(rows=2, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    col_widths = [Inches(3.4), Inches(3.4)]
    meta_data = [
        [("Project Ref:", " HUENICS-SYS-2026-REV2"), ("Date:", " August 21, 2026")],
        [("Client:", " Huenics Industrial Supply Corp."), ("Status:", " Final Approved Architecture & Scope")]
    ]
    
    for row_idx, row in enumerate(meta_table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = col_widths[col_idx]
            set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            lbl, val = meta_data[row_idx][col_idx]
            r1 = p.add_run(lbl)
            r1.font.bold = True
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            r2 = p.add_run(val)
            r2.font.size = Pt(9.5)
            r2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ─── SECTION 1: BUDGET & SCOPE ALLOCATION ─────────────────────────────────
    h1 = doc.add_heading(level=1)
    r_h1 = h1.add_run("1. Budget & Scope Comparison (Adjusted System Scope)")
    r_h1.font.size = Pt(13)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    p_intro = doc.add_paragraph(
        "The table below details the scope allocation under the optimized system architecture, incorporating all core backoffice ERP functions, OCR document ingestion, delivery lifecycle, free built-in warranty notifications, and the newly developed Public Customer Portal & Unofficial Quotation Generator:"
    )
    p_intro.paragraph_format.space_after = Pt(8)
    
    # Scope Comparison Table
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    headers = ["Functional Domain", "Base Proposed Scope", "Expanded / Newly Added Scope", "Scope Allocation & Status"]
    widths = [Inches(1.8), Inches(1.8), Inches(2.0), Inches(1.2)]
    
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        hdr_cells[i].width = widths[i]
        set_cell_background(hdr_cells[i], "0F172A")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=100, right=100)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    scope_rows = [
        ("PDF & Image Document Ingestion Engine", "Multi-vendor DB regex parser + PDF extraction", "+ High-Res Image OCR (JPG/PNG) + Lanczos upscaling + Dual-engine OCR (Windows Media OCR + Tesseract)", "100% INCLUDED\n(Visual Canvas deferred)"),
        ("Quotation & 1-Click PO Conversion", "Custom Quotes + 1-Click PO conversion + 12% VAT & line math audit", "Sales agent auto-crediting + Multi-tier approval (Draft → Reviewed → Approved) + E-signatures", "100% INCLUDED\n(Core Retained)"),
        ("🆕 Customer-Facing Storefront & Catalog", "Not in original scope", "• Corporate About Us profile\n• Searchable Product Catalog with category filtering\n• Real-time stock & price visibility", "100% INCLUDED\n(Newly Delivered)"),
        ("🆕 Online Quotation Builder (Customer)", "Not in original scope", "• Interactive BOQ Cart Builder with local persistence\n• Custom line item additions\n• Real-time 12% PH VAT math", "100% INCLUDED\n(Newly Delivered)"),
        ("🆕 Instant Unofficial Quotation PDF Engine", "Not in original scope", "• Automated UNOFF-YYYYMMDD-XXXX ref generator\n• Downloadable/Printable branded PDF estimate\n• Preliminary disclaimer & terms", "100% INCLUDED\n(Newly Delivered)"),
        ("Complete Delivery & Accounting Artifacts", "DR# logging & delivery tracking", "• Official Delivery Receipt (DR) PDF Generation\n• Official Sales Invoice (SI) PDF Generation with BIR numbering", "100% INCLUDED\n(Fully Integrated)"),
        ("Delivery & Warranty Tracker (Free Alerts)", "DR# logging, 6mo / 1yr / 2yrs warranty clocks", "Built-in 30-day warranty expiration in-app alerts & lifecycle tracking (100% Free / Core Feature)", "100% INCLUDED (FREE)\n(Core Retained)"),
        ("Modular BOM & Inventory Auto-Deduct", "1-tier modular BOM checklist (COB, Driver, Heatsink) + Auto-deduct", "Auto-deduction upon marking delivery as delivered + scaffold warehouse tracking", "100% INCLUDED\n(Streamlined 1-Tier)"),
        ("Sales Quota & Performance Analytics", "Real-time Leaderboard + Quota gauges + Realized Gross Profit", "Leaderboard rankings + Monthly win rate % + Real-time quota achievement gauges", "100% INCLUDED\n(Full Value)"),
        ("Multi-Auth Roles & Security Vault", "4-Role RBAC (Admin, Ops, Sales, CEO) + Private storage disk", "SHA-256 duplicate detection + Action Groups positioned before columns across all tables", "100% INCLUDED\n(Enterprise Security)"),
        ("Cloud & DevOps Infrastructure", "Local development setup", "• Multi-stage Docker containerization (PHP 8.3 + Node 22 + Nginx + Python)\n• Vercel Serverless zero-config deployment + 30/30 green tests", "100% INCLUDED\n(Production Ready)")
    ]
    
    for row_idx, data in enumerate(scope_rows):
        row = table.add_row()
        fill_color = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            cell = row.cells[col_idx]
            cell.width = widths[col_idx]
            set_cell_background(cell, fill_color)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            r.font.size = Pt(8.5)
            if col_idx == 0:
                r.font.bold = True
            elif col_idx == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
                
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Financial Summary Box Table
    fin_table = doc.add_table(rows=5, cols=2)
    fin_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fin_table.autofit = False
    f_widths = [Inches(4.8), Inches(2.0)]
    
    fin_data = [
        ("Original Backoffice System Baseline (PHP 63k Scope):", "PHP 63,000.00"),
        ("Strategic Backoffice Scope Optimizations (Deferred to Phase 2):", "- PHP 18,000.00"),
        ("Optimized Core Backoffice Baseline Budget:", "PHP 45,000.00"),
        ("Total Value of Newly Added Modules (Customer Portal, Unofficial Quote PDF, DR/SI):", "PHP 55,000.00 (VALUE)"),
        ("FINAL ADJUSTED TOTAL ALL-IN INVESTMENT (Comprehensive Package):", "PHP 50,000.00 NET")
    ]
    
    for r_idx, (lbl, val) in enumerate(fin_data):
        row = fin_table.rows[r_idx]
        is_total = (r_idx == 4)
        bg_col = "0F172A" if is_total else ("F1F5F9" if r_idx % 2 == 1 else "FFFFFF")
        
        for c_idx, cell in enumerate(row.cells):
            cell.width = f_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=90, bottom=90, left=120, right=120)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            text = lbl if c_idx == 0 else val
            r = p.add_run(text)
            if is_total:
                r.font.bold = True
                r.font.size = Pt(10.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if c_idx == 0 else RGBColor(0x38, 0xBD, 0xF8)
                if c_idx == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                r.font.size = Pt(9.5)
                if c_idx == 0:
                    r.font.bold = True
                if c_idx == 1:
                    r.font.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    if "-" in text:
                        r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
                        
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ─── SECTION 2: JUSTIFICATION FOR DEFERRED SCOPE ──────────────────────────
    h2 = doc.add_heading(level=1)
    r_h2 = h2.add_run("2. Detailed Justification for Deferred Scope (₱18,000 Total Savings)")
    r_h2.font.size = Pt(13)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    justifications = [
        ("1. Visual Coordinate Drag-and-Drop Layout GUI Builder (Savings: PHP 12,000.00)",
         "Why it can be deferred: The database-driven layout engine (vendor_document_layouts and vendor_layout_field_mappings) already processes unlimited vendor templates without code changes. System Administrators can manage vendor layout parameters via standard Filament forms or database seeders without needing an HTML5 canvas visual coordinate drag-and-drop tool.\nOperational Impact: ZERO. The parser extracts quote/PO numbers, line items, and totals automatically across all vendor templates."),
        
        ("2. Recursive Multi-Tier Sub-Assembly Hierarchy (Savings: PHP 6,000.00)",
         "Why it can be deferred: Huenics lighting assemblies (e.g., LED Tracklight MLTR 30W) utilize single-tier modular BOM components: COB (3000k/4000k), Driver (700mA/1050mA), and Heatsink. The streamlined modular component checklist satisfies 100% of actual assembly and auto-deduction requirements without needing deep recursive multi-level parent-child assembly trees.\nOperational Impact: ZERO. All sub-components are accurately auto-deducted from inventory upon marking delivery as delivered.")
    ]
    
    for title, desc in justifications:
        p_t = doc.add_paragraph()
        p_t.paragraph_format.space_after = Pt(2)
        r_t = p_t.add_run(title)
        r_t.font.bold = True
        r_t.font.size = Pt(10.5)
        r_t.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        
        p_d = doc.add_paragraph(desc)
        p_d.paragraph_format.space_after = Pt(6)
        p_d.paragraph_format.left_indent = Inches(0.2)
        p_d.style.font.size = Pt(9.5)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # ─── SECTION 3: COMPLETE GUARANTEED DELIVERABLES ──────────────────────────
    h3 = doc.add_heading(level=1)
    r_h3 = h3.add_run("3. Complete Guaranteed Deliverables (Comprehensive System Scope)")
    r_h3.font.size = Pt(13)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    deliverable_categories = [
        ("A. Public Customer Portal & Online Quotation Generator (NEW)", [
            ("About Huenics Profile (/about):", " Corporate history, product lines (Plumbing, Structural Steel, Pumps, Electrical), and BIR tax credentials."),
            ("Product Showcase & Catalog (/products):", " Interactive searchable catalog with category filtering, SKU badges, unit prices, and instant Add to Quote controls."),
            ("Interactive Quotation / BOQ Builder (/quotation/builder):", " Live cart builder with browser localStorage persistence, custom non-catalog line items insertion, and real-time line calculations."),
            ("Automated 12% Philippine VAT Calculation:", " Real-time computation of Net Vatable Subtotal, 12% VAT, and Grand Total."),
            ("Instant Unofficial Quotation PDF Generator:", " Formatted commercial estimate (UNOFF-YYYYMMDD-XXXX) with official header, itemized table, VAT breakdown, and preliminary disclaimer.")
        ]),
        ("B. Document Ingestion & Computer Vision OCR Engine", [
            ("Multi-Format Extraction (PDF & Image):", " Ingests PDF documents as well as high-res images (JPG/PNG) of invoices and POs."),
            ("Computer Vision Preprocessing:", " Lanczos upscaling, grayscale binarization, adaptive contrast enhancement, and noise reduction."),
            ("Dual OCR Engine Integration:", " Windows Media OCR for local Windows systems + Linux native Tesseract fallback."),
            ("12% Philippine VAT & Line Math Verification:", " Automated detection and red-flagging of vendor arithmetic errors (e.g., .85 line error, ₱112,500.00 copied VAT discrepancy)."),
            ("Private Document Storage Vault:", " SHA-256 duplicate upload prevention and protected private storage disk with strict authentication middleware.")
        ]),
        ("C. Sales, Quotations & Purchase Order Management", [
            ("End-to-End Quotation Management:", " Custom customer quotes with auto-numbering (QT-YYYY-NNNN), line items with real-time profit computation, and multi-tier approval status."),
            ("1-Click Quotation ➔ Purchase Order Conversion:", " 1-Click conversion copying customer details, line items, auto-computing 12% PH VAT, and crediting the sales agent's monthly sales quota."),
            ("Customer E-Signature & Verification:", " E-signature capture and verification flag for conversion into official POs.")
        ]),
        ("D. Delivery, Invoicing, Warranty (Free Alerts) & Inventory Auto-Deduct", [
            ("Delivery Receipt (DR) Lifecycle & PDF Export:", " Dedicated DR generation (DR-YYYY-XXXX), dispatch vehicle/driver logging, site receiver verification, and printable PDF export."),
            ("Sales Invoice (SI) Lifecycle & PDF Export:", " BIR-compliant Sales Invoices (SI-YYYY-XXXX) with payment status tracking (Unpaid, Partially Paid, Paid) and formal PDF export."),
            ("Warranty Lifecycle Tracking (Free Built-In Alerts):", " Delivery date-anchored warranty clocks with 3 fixed standard options: 1 Year (1yr), 2 Years (2yrs), and 6 Months, complete with automated 30-day expiration notifications (100% Free / Core Feature)."),
            ("Modular Component BOM Auto-Deduction:", " Checklists for COB, Driver, and Heatsink parts auto-deducting from stock when the PO is marked delivered.")
        ]),
        ("E. Executive Analytics, Security & DevOps", [
            ("Full Real-Time Sales Leaderboard & Quota Gauges:", " Live ranked sales leaderboard, monthly target achievement gauges (₱ achieved vs target), win rate %, and realized gross profit engine."),
            ("Multi-Auth Role-Based Access Control (RBAC):", " 4 distinct secure user roles: Admin, Operations Manager, Sales Executive, and CEO."),
            ("Action Group UI Optimization:", " All row actions consolidated in clean dropdown ActionGroup elements positioned before table columns."),
            ("Multi-Stage Docker & Vercel Cloud Deployment:", " Multi-stage production container setup (Dockerfile, docker-compose.yml) and zero-downtime Vercel serverless deployment with automated SQLite fallback."),
            ("Automated Test Suite:", " 30 comprehensive unit & feature test cases passing 100% green.")
        ])
    ]
    
    for cat_title, items in deliverable_categories:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_before = Pt(4)
        p_c.paragraph_format.space_after = Pt(2)
        r_c = p_c.add_run(cat_title)
        r_c.font.bold = True
        r_c.font.size = Pt(11)
        r_c.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        for item_lbl, item_desc in items:
            p_i = doc.add_paragraph(style='List Bullet')
            p_i.paragraph_format.space_after = Pt(2)
            r_il = p_i.add_run(item_lbl)
            r_il.font.bold = True
            r_il.font.size = Pt(9.5)
            r_id = p_i.add_run(item_desc)
            r_id.font.size = Pt(9.5)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    
    # ─── SECTION 4: PAYMENT TERMS & MILESTONE SCHEDULE ────────────────────────
    h4 = doc.add_heading(level=1)
    r_h4 = h4.add_run("4. Payment Terms & Milestone Schedule")
    r_h4.font.size = Pt(13)
    r_h4.font.bold = True
    r_h4.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    
    m_table = doc.add_table(rows=1, cols=4)
    m_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_table.autofit = False
    
    m_headers = ["Milestone Stage", "Key Deliverables Included", "Percentage", "Amount (PHP)"]
    m_widths = [Inches(1.8), Inches(3.2), Inches(0.9), Inches(1.1)]
    
    m_hdr_cells = m_table.rows[0].cells
    for i, h_text in enumerate(m_headers):
        m_hdr_cells[i].width = m_widths[i]
        set_cell_background(m_hdr_cells[i], "0F172A")
        set_cell_margins(m_hdr_cells[i], top=100, bottom=100, left=100, right=100)
        p = m_hdr_cells[i].paragraphs[0]
        r = p.add_run(h_text)
        r.font.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        
    m_rows = [
        ("Milestone 1 (Mobilization)", "System Architecture, Database Schema, 4-Role RBAC & Customer Storefront Portal", "30%", "PHP 15,000.00"),
        ("Milestone 2 (Core Engine)", "OCR Document Pipeline, VAT Math Engine, Quotation → PO → DR → SI Full Lifecycle", "40%", "PHP 20,000.00"),
        ("Milestone 3 (Final Delivery)", "Unofficial Quotation PDF Engine, Docker Setup, Vercel Serverless Deployment & Handover", "30%", "PHP 15,000.00"),
        ("TOTAL INVESTMENT", "Comprehensive All-In Enterprise Package (All Deliverables Included)", "100%", "PHP 50,000.00 NET")
    ]
    
    for r_idx, m_data in enumerate(m_rows):
        row = m_table.add_row()
        is_tot = (r_idx == 3)
        bg_col = "0F172A" if is_tot else ("F8FAFC" if r_idx % 2 == 1 else "FFFFFF")
        for c_idx, text in enumerate(m_data):
            cell = row.cells[c_idx]
            cell.width = m_widths[c_idx]
            set_cell_background(cell, bg_col)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            if is_tot:
                r.font.bold = True
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if c_idx < 3 else RGBColor(0x38, 0xBD, 0xF8)
            else:
                r.font.size = Pt(9)
                if c_idx == 0 or c_idx == 3:
                    r.font.bold = True
                    
    # Sign-off block
    doc.add_paragraph().paragraph_format.space_after = Pt(16)
    
    sig_table = doc.add_table(rows=2, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_table.autofit = False
    
    for row in sig_table.rows:
        for cell in row.cells:
            cell.width = Inches(3.4)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            
    p1 = sig_table.rows[0].cells[0].paragraphs[0]
    p1.add_run("Prepared & Submitted by:\n\n\n_________________________________\nLead Solutions Architect & Engineer\nDate: August 21, 2026").font.size = Pt(9)
    
    p2 = sig_table.rows[0].cells[1].paragraphs[0]
    p2.add_run("Conforme & Approved by:\n\n\n_________________________________\nAuthorized Executive / Client Representative\nDate: ________________________").font.size = Pt(9)
    
    # Save document
    output_path = r"c:\laragon\www\huenics\HUENICS_PROPOSAL_ADJUSTED.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    create_document()
